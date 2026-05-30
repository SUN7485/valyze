"""
Export Service - Multi-format export for credit reports.
Supports: JSON, XML, Excel (XLSX), CSV, Word (DOCX)
"""

from __future__ import annotations

import csv
import io
import json
import os
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from xml.etree.ElementTree import Element, SubElement, ElementTree, tostring


OUTPUT_DIR = Path("/tmp/outputs")

try:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
except OSError:
    pass  # Vercel serverless - directory will be created on first request


def _ensure_output_dir():
    """Create output directory if it doesn't exist (lazy init for Vercel serverless)."""
    try:
        OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    except OSError:
        pass


def _get_company_name(report: Dict[str, Any]) -> str:
    """Extract company name from report."""
    fields = report.get("fields", {})
    company_name = fields.get("company_name", {}).get("value")
    if not company_name:
        company_name = fields.get("legal_name", {}).get("value")
    return str(company_name or "Report").replace(" ", "_")[:30]


def _safe_value(val: Any) -> str:
    """Convert value to safe string."""
    if val is None:
        return ""
    if isinstance(val, dict):
        return val.get("value", "") or ""
    return str(val)


def _get_field_value(fields: Dict, key: str) -> Any:
    """Get field value from fields dict."""
    field = fields.get(key)
    if field is None:
        return None
    if isinstance(field, dict):
        return field.get("value")
    return field


def _export_json(report: Dict[str, Any], report_id: str) -> Path:
    """Export report as JSON."""
    _ensure_output_dir()
    output_path = OUTPUT_DIR / f"{report_id}.json"

    # Convert report to clean JSON structure
    clean_report = {
        "report_id": report.get("report_id"),
        "status": report.get("status"),
        "created_at": report.get("created_at"),
        "updated_at": report.get("updated_at"),
        "fields": {},
        "arrays": {},
        "files": report.get("files", []),
        "extraction_stats": report.get("extraction_stats", {}),
    }

    # Clean fields
    fields = report.get("fields", {})
    for key, value in fields.items():
        if isinstance(value, dict):
            clean_report["fields"][key] = value.get("value")
        else:
            clean_report["fields"][key] = value

    # Clean arrays
    arrays = report.get("arrays", {})
    for key, value in arrays.items():
        if isinstance(value, list):
            clean_report["arrays"][key] = value
        else:
            clean_report["arrays"][key] = []

    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(clean_report, f, indent=2, ensure_ascii=False)

    return output_path


def _export_xml(report: Dict[str, Any], report_id: str) -> Path:
    """Export report as XML."""
    _ensure_output_dir()
    output_path = OUTPUT_DIR / f"{report_id}.xml"

    root = Element("Report")
    root.set("report_id", str(report.get("report_id", "")))
    root.set("status", str(report.get("status", "")))
    root.set("created_at", str(report.get("created_at", "")))

    # Fields section
    fields_elem = SubElement(root, "Fields")
    fields = report.get("fields", {})
    for key, value in fields.items():
        field_elem = SubElement(fields_elem, "Field")
        field_elem.set("name", key)

        if isinstance(value, dict):
            val = value.get("value")
            conf = value.get("confidence", "")
            src = value.get("source", "")
            if val is not None:
                field_elem.set("value", str(val))
            if conf:
                field_elem.set("confidence", conf)
            if src:
                field_elem.set("source", src)
        else:
            if value is not None:
                field_elem.set("value", str(value))

    # Arrays section
    arrays_elem = SubElement(root, "Arrays")
    arrays = report.get("arrays", {})
    for array_name, items in arrays.items():
        array_elem = SubElement(arrays_elem, array_name)
        if isinstance(items, list):
            for idx, item in enumerate(items):
                item_elem = SubElement(array_elem, "Item")
                item_elem.set("index", str(idx))
                if isinstance(item, dict):
                    for k, v in item.items():
                        if v is not None:
                            sub_elem = SubElement(item_elem, k)
                            sub_elem.text = str(v)
                elif item is not None:
                    item_elem.text = str(item)

    # Files section
    files_elem = SubElement(root, "Files")
    for file_info in report.get("files", []):
        if isinstance(file_info, dict):
            file_elem = SubElement(files_elem, "File")
            for k, v in file_info.items():
                if v is not None:
                    file_elem.set(k, str(v))

    xml_str = tostring(root, encoding="unicode", method="xml")
    xml_str = '<?xml version="1.0" encoding="UTF-8"?>\n' + xml_str

    with open(output_path, "w", encoding="utf-8") as f:
        f.write(xml_str)

    return output_path


def _export_excel(report: Dict[str, Any], report_id: str) -> Path:
    """Export report as Excel workbook."""
    _ensure_output_dir()
    output_path = OUTPUT_DIR / f"{report_id}.xlsx"

    wb = Workbook()
    wb.remove(wb.active)

    # Styles
    header_font = Font(bold=True, size=12)
    header_fill = PatternFill(
        start_color="4472C4", end_color="4472C4", fill_type="solid"
    )
    header_font_white = Font(bold=True, size=12, color="FFFFFF")

    # Sheet 1: Summary
    ws_summary = wb.create_sheet("Summary", 0)
    ws_summary.append(["Field", "Value"])
    ws_summary["A1"].font = header_font_white
    ws_summary["B1"].font = header_font_white
    ws_summary["A1"].fill = header_fill
    ws_summary["B1"].fill = header_fill

    fields = report.get("fields", {})
    important_fields = [
        "company_name",
        "legal_name",
        "cr_number",
        "country",
        "city",
        "industry",
        "business_summary",
        "established_year",
        "employee_count",
        "revenue",
        "net_profit",
        "total_assets",
        "total_equity",
    ]

    for key in important_fields:
        val = _get_field_value(fields, key)
        if val:
            ws_summary.append([key.replace("_", " ").title(), _safe_value(val)])

    for cell in ws_summary[1]:
        cell.font = header_font

    # Sheet 2: All Fields
    ws_fields = wb.create_sheet("All Fields", 1)
    ws_fields.append(["Field Name", "Value", "Confidence", "Source"])
    ws_fields["A1"].font = header_font_white
    ws_fields["B1"].font = header_font_white
    ws_fields["C1"].font = header_font_white
    ws_fields["D1"].font = header_font_white
    ws_fields["A1"].fill = header_fill
    ws_fields["B1"].fill = header_fill
    ws_fields["C1"].fill = header_fill
    ws_fields["D1"].fill = header_fill

    for key, value in fields.items():
        if isinstance(value, dict):
            val = value.get("value", "")
            conf = value.get("confidence", "")
            src = value.get("source", "")
            ws_fields.append([key, _safe_value(val), conf, src])
        else:
            ws_fields.append([key, _safe_value(value), "", ""])

    for cell in ws_fields[1]:
        cell.font = header_font

    # Sheet 3: Shareholders
    ws_shares = wb.create_sheet("Shareholders", 2)
    ws_shares.append(["Name", "Percentage", "Nationality", "Type"])
    ws_shares["A1"].font = header_font_white
    ws_shares["B1"].font = header_font_white
    ws_shares["C1"].font = header_font_white
    ws_shares["D1"].font = header_font_white
    for cell in ws_shares[1]:
        cell.fill = header_fill

    shareholders = report.get("arrays", {}).get("shareholders", [])
    for sh in shareholders:
        if isinstance(sh, dict):
            ws_shares.append(
                [
                    sh.get("name", ""),
                    _safe_value(sh.get("percentage") or sh.get("ownership_percentage")),
                    sh.get("nationality", ""),
                    sh.get("type", ""),
                ]
            )

    for cell in ws_shares[1]:
        cell.font = header_font

    # Sheet 4: Branches
    ws_branches = wb.create_sheet("Branches", 3)
    ws_branches.append(["Branch Name", "City", "Status", "CR Number"])
    ws_branches["A1"].font = header_font_white
    ws_branches["B1"].font = header_font_white
    ws_branches["C1"].font = header_font_white
    ws_branches["D1"].font = header_font_white
    for cell in ws_branches[1]:
        cell.fill = header_fill

    branches = report.get("arrays", {}).get("branches", [])
    for br in branches:
        if isinstance(br, dict):
            ws_branches.append(
                [
                    br.get("branch_name", ""),
                    br.get("branch_city", ""),
                    br.get("branch_status", ""),
                    br.get("branch_cr_no", ""),
                ]
            )

    for cell in ws_branches[1]:
        cell.font = header_font

    # Sheet 5: Banking
    ws_banking = wb.create_sheet("Banking", 4)
    ws_banking.append(["Bank Name", "Facility Type", "Facility Usage"])
    ws_banking["A1"].font = header_font_white
    ws_banking["B1"].font = header_font_white
    ws_banking["C1"].font = header_font_white
    for cell in ws_banking[1]:
        cell.fill = header_fill

    banking = report.get("arrays", {}).get("banking_relationships", [])
    for b in banking:
        if isinstance(b, dict):
            ws_banking.append(
                [
                    b.get("bank_name", ""),
                    b.get("facility_type", ""),
                    b.get("facility_usage", ""),
                ]
            )

    for cell in ws_banking[1]:
        cell.font = header_font

    # Sheet 6: Management Team
    ws_team = wb.create_sheet("Management Team", 5)
    ws_team.append(["Name", "Title", "Department", "Email"])
    ws_team["A1"].font = header_font_white
    ws_team["B1"].font = header_font_white
    ws_team["C1"].font = header_font_white
    ws_team["D1"].font = header_font_white
    for cell in ws_team[1]:
        cell.fill = header_fill

    team = report.get("arrays", {}).get("management_team", [])
    for m in team:
        if isinstance(m, dict):
            ws_team.append(
                [
                    m.get("name", ""),
                    m.get("title", ""),
                    m.get("department", ""),
                    m.get("contact_email", ""),
                ]
            )

    for cell in ws_team[1]:
        cell.font = header_font

    # Adjust column widths
    for ws in wb.worksheets:
        for col in ws.columns:
            max_length = 0
            col_letter = col[0].column_letter
            for cell in col:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)
            ws.column_dimensions[col_letter].width = adjusted_width

    wb.save(output_path)
    return output_path


def _export_csv(report: Dict[str, Any], report_id: str) -> Path:
    """Export report as CSV (main fields only)."""
    _ensure_output_dir()
    output_path = OUTPUT_DIR / f"{report_id}.csv"

    fields = report.get("fields", {})
    important_fields = [
        "company_name",
        "legal_name",
        "cr_number",
        "tax_registration_number",
        "country",
        "city",
        "address",
        "industry",
        "business_summary",
        "established_year",
        "employee_count",
        "legal_form",
        "revenue",
        "net_profit",
        "total_assets",
        "total_equity",
        "total_liabilities",
        "contact_email",
        "contact_phone",
    ]

    rows = [["Field", "Value"]]
    for key in important_fields:
        val = _get_field_value(fields, key)
        if val:
            rows.append([key.replace("_", " ").title(), _safe_value(val)])

    with open(output_path, "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerows(rows)

    return output_path


def _export_word(report: Dict[str, Any], report_id: str) -> Path:
    """Export report as Word document."""
    _ensure_output_dir()
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.table import Wd_Table_Alignment
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        # python-docx not installed, create simple text file
        output_path = OUTPUT_DIR / f"{report_id}.txt"
        company_name = _get_company_name(report)

        with open(output_path, "w", encoding="utf-8") as f:
            f.write(f"CREDIT REPORT: {company_name}\n")
            f.write("=" * 50 + "\n\n")

            fields = report.get("fields", {})
            for key, value in fields.items():
                val = _get_field_value(fields, key)
                if val:
                    f.write(f"{key.replace('_', ' ').title()}: {_safe_value(val)}\n")

            f.write("\n" + "=" * 50 + "\n")
            f.write("SHAREHOLDERS\n")
            f.write("=" * 50 + "\n")
            for sh in report.get("arrays", {}).get("shareholders", []):
                if isinstance(sh, dict):
                    f.write(
                        f"- {sh.get('name')}: {sh.get('percentage') or sh.get('ownership_percentage')}\n"
                    )

        return output_path

    output_path = OUTPUT_DIR / f"{report_id}.docx"
    doc = Document()

    # Title
    company_name = _get_company_name(report)
    title = doc.add_heading(f"CREDIT REPORT: {company_name}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Summary section
    doc.add_heading("Company Information", level=1)
    fields = report.get("fields", {})

    important_fields = [
        "company_name",
        "legal_name",
        "cr_number",
        "tax_registration_number",
        "country",
        "city",
        "address",
        "industry",
        "business_summary",
        "established_year",
        "employee_count",
        "legal_form",
    ]

    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"

    for key in important_fields:
        val = _get_field_value(fields, key)
        if val:
            row = table.add_row()
            row.cells[0].text = key.replace("_", " ").title()
            row.cells[1].text = _safe_value(val)

    # Financial Summary
    doc.add_heading("Financial Summary", level=1)
    financial_fields = [
        "revenue",
        "net_profit",
        "total_assets",
        "total_equity",
        "total_liabilities",
    ]

    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"

    for key in financial_fields:
        val = _get_field_value(fields, key)
        if val:
            row = table.add_row()
            row.cells[0].text = key.replace("_", " ").title()
            row.cells[1].text = _safe_value(val)

    # Shareholders
    shareholders = report.get("arrays", {}).get("shareholders", [])
    if shareholders:
        doc.add_heading("Shareholders", level=1)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        table.rows[0].cells[0].text = "Name"
        table.rows[0].cells[1].text = "Percentage"
        table.rows[0].cells[2].text = "Nationality"

        for sh in shareholders:
            if isinstance(sh, dict):
                row = table.add_row()
                row.cells[0].text = sh.get("name", "")
                row.cells[1].text = _safe_value(
                    sh.get("percentage") or sh.get("ownership_percentage")
                )
                row.cells[2].text = sh.get("nationality", "")

    # Branches
    branches = report.get("arrays", {}).get("branches", [])
    if branches:
        doc.add_heading("Branches", level=1)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        table.rows[0].cells[0].text = "Branch Name"
        table.rows[0].cells[1].text = "City"
        table.rows[0].cells[2].text = "Status"

        for br in branches:
            if isinstance(br, dict):
                row = table.add_row()
                row.cells[0].text = br.get("branch_name", "")
                row.cells[1].text = br.get("branch_city", "")
                row.cells[2].text = br.get("branch_status", "")

    # Banking Relationships
    banking = report.get("arrays", {}).get("banking_relationships", [])
    if banking:
        doc.add_heading("Banking Relationships", level=1)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        table.rows[0].cells[0].text = "Bank Name"
        table.rows[0].cells[1].text = "Facility Type"
        table.rows[0].cells[2].text = "Facility Usage"

        for b in banking:
            if isinstance(b, dict):
                row = table.add_row()
                row.cells[0].text = b.get("bank_name", "")
                row.cells[1].text = b.get("facility_type", "")
                row.cells[2].text = b.get("facility_usage", "")

    # Management Team
    team = report.get("arrays", {}).get("management_team", [])
    if team:
        doc.add_heading("Management Team", level=1)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        table.rows[0].cells[0].text = "Name"
        table.rows[0].cells[1].text = "Title"
        table.rows[0].cells[2].text = "Department"

        for m in team:
            if isinstance(m, dict):
                row = table.add_row()
                row.cells[0].text = m.get("name", "")
                row.cells[1].text = m.get("title", "")
                row.cells[2].text = m.get("department", "")

    doc.save(output_path)
    return output_path


# ---------------------------------------------------------------------------
# Main export functions
# ---------------------------------------------------------------------------


def export_json(report: Dict[str, Any], report_id: str) -> Dict[str, Any]:
    """Export report as JSON."""
    try:
        output_path = _export_json(report, report_id)
        return {
            "success": True,
            "format": "json",
            "file_path": str(output_path),
            "file_size_kb": round(output_path.stat().st_size / 1024, 1),
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


def export_xml(report: Dict[str, Any], report_id: str) -> Dict[str, Any]:
    """Export report as XML."""
    try:
        output_path = _export_xml(report, report_id)
        return {
            "success": True,
            "format": "xml",
            "file_path": str(output_path),
            "file_size_kb": round(output_path.stat().st_size / 1024, 1),
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


def export_excel(report: Dict[str, Any], report_id: str) -> Dict[str, Any]:
    """Export report as Excel."""
    try:
        output_path = _export_excel(report, report_id)
        return {
            "success": True,
            "format": "xlsx",
            "file_path": str(output_path),
            "file_size_kb": round(output_path.stat().st_size / 1024, 1),
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


def export_csv(report: Dict[str, Any], report_id: str) -> Dict[str, Any]:
    """Export report as CSV."""
    try:
        output_path = _export_csv(report, report_id)
        return {
            "success": True,
            "format": "csv",
            "file_path": str(output_path),
            "file_size_kb": round(output_path.stat().st_size / 1024, 1),
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


def export_word(report: Dict[str, Any], report_id: str) -> Dict[str, Any]:
    """Export report as Word."""
    try:
        output_path = _export_word(report, report_id)
        return {
            "success": True,
            "format": "docx",
            "file_path": str(output_path),
            "file_size_kb": round(output_path.stat().st_size / 1024, 1),
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


def get_filename(report: Dict[str, Any], format: str) -> str:
    """Generate safe filename for report."""
    company_name = _get_company_name(report)
    return f"CreditReport_{company_name}.{format}"


def check_export_files(report_id: str) -> Dict[str, Any]:
    """Check which export files exist."""
    formats = ["json", "xml", "xlsx", "csv", "docx", "pdf"]
    result = {}

    for fmt in formats:
        path = OUTPUT_DIR / f"{report_id}.{fmt}"
        if path.exists():
            result[fmt] = {
                "exists": True,
                "size_kb": round(path.stat().st_size / 1024, 1),
            }
        else:
            result[fmt] = {"exists": False}

    return result
