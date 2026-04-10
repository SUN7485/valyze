"""
Word Document Extractor forValyze Credit report.

Handles .docx files via python-docx. Provides best-effort
support for legacy .doc files.
"""

from __future__ import annotations

from pathlib import Path
from typing import Optional


class WordExtractor:
    """Extracts text and tables from Word documents."""

    # ------------------------------------------------------------------
    # Main entry point
    # ------------------------------------------------------------------

    def extract(self, file_path: Path) -> dict:
        """Extract text and tables from a Word document."""
        result = {
            "text": "",
            "pages": None,
            "language": "english",
            "file_type": "word",
            "tables": [],
            "success": False,
            "error": None,
        }

        ext = file_path.suffix.lower()

        try:
            if ext == ".docx":
                print(f"[WORD] Processing DOCX: {file_path.name}")
                text = self.extract_text_from_docx(file_path)
                tables = self.extract_tables_from_docx(file_path)
                result["text"] = text
                result["tables"] = tables
                result["success"] = bool(text.strip())
                print(f"[WORD] Extracted {len(text)} chars, {len(tables)} tables")
            elif ext == ".doc":
                print(f"[WORD] Processing legacy DOC: {file_path.name}")
                text = self.handle_doc_file(file_path)
                result["text"] = text
                result["success"] = bool(text.strip())
            else:
                result["error"] = f"Unsupported Word format: {ext}"

            if not result["text"].strip():
                result["error"] = result["error"] or "No text extracted from Word document"

        except Exception as e:
            result["error"] = f"Word extraction failed: {e}"
            print(f"[WORD] ERROR: {e}")

        return result

    # ------------------------------------------------------------------
    # DOCX text extraction
    # ------------------------------------------------------------------

    def extract_text_from_docx(self, file_path: Path) -> str:
        """Extract all text from a DOCX file (paragraphs + table cells)."""
        try:
            from docx import Document

            doc = Document(str(file_path))
            parts: list[str] = []

            # Extract paragraphs (preserves order, handles RTL Arabic)
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    parts.append(text)

            # Extract text from table cells
            for table in doc.tables:
                for row in table.rows:
                    row_texts = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_texts.append(cell_text)
                    if row_texts:
                        parts.append(" | ".join(row_texts))

            return "\n".join(parts)

        except ImportError:
            print("[WORD] python-docx not installed")
            return ""
        except Exception as e:
            print(f"[WORD] DOCX text extraction error: {e}")
            return ""

    # ------------------------------------------------------------------
    # DOCX table extraction
    # ------------------------------------------------------------------

    def extract_tables_from_docx(self, file_path: Path) -> list:
        """Extract tables from a DOCX file in the standard table format."""
        tables: list[dict] = []

        try:
            from docx import Document

            doc = Document(str(file_path))

            for table in doc.tables:
                if len(table.rows) < 2:
                    continue

                raw: list[list[str]] = []
                for row in table.rows:
                    cells = [
                        (cell.text.strip() if cell.text else "")
                        for cell in row.cells
                    ]
                    raw.append(cells)

                if len(raw[0]) < 2:
                    continue

                # Check if first row looks like headers
                headers = raw[0]
                rows = raw[1:]

                tables.append({
                    "page": None,  # Word doesn't have page numbers per table
                    "headers": headers,
                    "rows": rows,
                    "raw": raw,
                })

        except ImportError:
            print("[WORD] python-docx not installed — skipping table extraction")
        except Exception as e:
            print(f"[WORD] Table extraction error: {e}")

        return tables

    # ------------------------------------------------------------------
    # Legacy .doc handling
    # ------------------------------------------------------------------

    def handle_doc_file(self, file_path: Path) -> str:
        """
        Attempt to read a legacy .doc file.

        python-docx does not support .doc format, so this is best-effort.
        """
        try:
            from docx import Document
            doc = Document(str(file_path))
            parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            if parts:
                return "\n".join(parts)
        except Exception:
            pass

        print(
            f"[WORD] Cannot read legacy .doc file: {file_path.name}\n"
            "  Please convert to .docx format for full extraction."
        )
        return (
            f"[Warning] File '{file_path.name}' is in legacy .doc format. "
            "Please convert to .docx for full text extraction."
        )
