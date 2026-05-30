"""
Export API - Multi-format export endpoints.
All generation is done inline to avoid import failures in Vercel serverless.
"""

from __future__ import annotations

import csv
import io
import json as json_mod
import base64
from datetime import datetime, timezone
from pathlib import Path
from xml.etree.ElementTree import Element, SubElement, tostring

from fastapi import APIRouter, HTTPException
from fastapi.responses import Response

from database.crud import get_report

router = APIRouter(prefix="/api/export", tags=["export"])


def _get_fields(report) -> dict:
    """Extract fields dict from report (handles both model and dict)."""
    fields = report.fields if hasattr(report, "fields") else report.get("fields", {})
    if hasattr(fields, "items"):
        return fields
    return {}


def _fv(fields: dict, key: str, default="") -> str:
    """Get a field value as string."""
    f = fields.get(key)
    if f is None:
        return default
    if hasattr(f, "value"):
        v = f.value
    elif isinstance(f, dict):
        v = f.get("value")
    else:
        v = f
    if v is None:
        return default
    return str(v)


def _safe(val) -> str:
    if val is None:
        return ""
    return str(val)


async def _get_report_or_404(report_id: str):
    report = await get_report(None, report_id)
    if report is None:
        raise HTTPException(404, f"Report {report_id} not found")
    return report


def _company_name(report) -> str:
    fields = _get_fields(report)
    name = _fv(fields, "company_name") or _fv(fields, "legal_name") or "Report"
    return name


# ---------------------------------------------------------------------------
# Export Endpoints — all return content directly
# ---------------------------------------------------------------------------


@router.post("/json/{report_id}")
async def export_report_json(report_id: str):
    """Export report as JSON file download."""
    report = await _get_report_or_404(report_id)
    data = report.model_dump() if hasattr(report, "model_dump") else report
    content = json_mod.dumps(data, indent=2, default=str, ensure_ascii=False)
    return Response(
        content=content,
        media_type="application/json",
        headers={"Content-Disposition": f'attachment; filename="{report_id}.json"'},
    )


@router.post("/xml/{report_id}")
async def export_report_xml(report_id: str):
    """Export report as XML file download."""
    report = await _get_report_or_404(report_id)
    fields = _get_fields(report)

    root = Element("CreditReport")
    for key in sorted(fields.keys()):
        val = _fv(fields, key)
        if val:
            SubElement(root, key).text = val

    xml_bytes = tostring(root, encoding="unicode", method="xml")
    xml_str = '<?xml version="1.0" encoding="UTF-8"?>\n' + xml_bytes

    return Response(
        content=xml_str,
        media_type="application/xml",
        headers={"Content-Disposition": f'attachment; filename="{report_id}.xml"'},
    )


@router.post("/csv/{report_id}")
async def export_report_csv(report_id: str):
    """Export report as CSV file download."""
    report = await _get_report_or_404(report_id)
    fields = _get_fields(report)

    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow(["Field", "Value"])
    for key in sorted(fields.keys()):
        writer.writerow([key, _fv(fields, key)])

    return Response(
        content=buf.getvalue(),
        media_type="text/csv",
        headers={"Content-Disposition": f'attachment; filename="{_company_name(report)}.csv"'},
    )


@router.post("/excel/{report_id}")
async def export_report_excel(report_id: str):
    """Export report as Excel — returns base64-encoded XLSX."""
    report = await _get_report_or_404(report_id)
    fields = _get_fields(report)

    try:
        from openpyxl import Workbook
        from openpyxl.styles import Alignment, Font, PatternFill

        wb = Workbook()
        ws = wb.active
        ws.title = "Credit Report"

        # Header style
        header_font = Font(bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="2563EB", end_color="2563EB", fill_type="solid")

        ws.cell(row=1, column=1, value="Field").font = header_font
        ws.cell(row=1, column=1).fill = header_fill
        ws.cell(row=1, column=2, value="Value").font = header_font
        ws.cell(row=1, column=2).fill = header_fill

        for i, key in enumerate(sorted(fields.keys()), start=2):
            ws.cell(row=i, column=1, value=key)
            ws.cell(row=i, column=2, value=_fv(fields, key))

        ws.column_dimensions["A"].width = 35
        ws.column_dimensions["B"].width = 60

        buf = io.BytesIO()
        wb.save(buf)
        buf.seek(0)
        b64 = base64.b64encode(buf.read()).decode()

        return {"success": True, "base64": b64, "filename": f"{_company_name(report)}.xlsx"}
    except ImportError:
        raise HTTPException(500, "Excel export requires openpyxl — not installed")
    except Exception as e:
        raise HTTPException(500, f"Excel export failed: {str(e)}")


@router.post("/word/{report_id}")
async def export_report_word(report_id: str):
    """Export report as Word — returns base64-encoded DOCX."""
    report = await _get_report_or_404(report_id)
    fields = _get_fields(report)

    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        doc.add_heading(_company_name(report), level=0)
        doc.add_heading("Credit Report", level=1)

        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Field"
        hdr[1].text = "Value"

        for key in sorted(fields.keys()):
            row = table.add_row().cells
            row[0].text = key
            row[1].text = _fv(fields, key)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        b64 = base64.b64encode(buf.read()).decode()

        return {"success": True, "base64": b64, "filename": f"{_company_name(report)}.docx"}
    except ImportError:
        raise HTTPException(500, "Word export requires python-docx — not installed")
    except Exception as e:
        raise HTTPException(500, f"Word export failed: {str(e)}")


# ---------------------------------------------------------------------------
# Legacy download endpoint (kept for backward compatibility)
# ---------------------------------------------------------------------------

OUTPUT_DIR = Path("/tmp/outputs")


@router.get("/download/{report_id}/{format}")
async def download_export(report_id: str, format: str):
    """Download previously generated export file."""
    filename = f"{report_id}.{format}"
    file_path = OUTPUT_DIR / filename
    if not file_path.exists():
        raise HTTPException(404, "Export file not found")
    from fastapi.responses import FileResponse
    return FileResponse(
        path=str(file_path),
        filename=filename,
        media_type="application/octet-stream",
    )


@router.get("/status/{report_id}")
async def export_status(report_id: str):
    """Check export generation status."""
    return {
        "report_id": report_id,
        "engine": "inline",
        "message": "Exports are generated on-demand",
    }


# ---------------------------------------------------------------------------
# Backup Endpoints
# ---------------------------------------------------------------------------


@router.get("/backup/all")
async def export_all_reports_backup():
    """Export all reports as JSON."""
    from database.crud import get_all_reports
    reports = await get_all_reports(None)
    if not reports:
        return {"message": "No reports to backup", "count": 0}
    return {
        "version": "1.0.0",
        "exported_at": datetime.now(timezone.utc).isoformat(),
        "total_reports": len(reports),
        "reports": [r.model_dump() if hasattr(r, "model_dump") else r for r in reports],
    }


@router.get("/backup/download/all")
async def download_all_reports_backup():
    """Download all reports as a downloadable JSON backup."""
    from database.crud import get_all_reports
    reports = await get_all_reports(None)
    if not reports:
        raise HTTPException(404, "No reports to backup")
    backup_data = {
        "version": "1.0.0",
        "exported_at": datetime.now(timezone.utc).isoformat(),
        "total_reports": len(reports),
        "reports": [r.model_dump() if hasattr(r, "model_dump") else r for r in reports],
    }
    content = json_mod.dumps(backup_data, indent=2, default=str)
    return Response(
        content=content,
        media_type="application/json",
        headers={"Content-Disposition": 'attachment; filename="valyze_backup.json"'},
    )